"""
docx_to_csv.py

Extracts text from all .docx files in a folder, chunks into 1-3 sentence blurbs,
applies rule-based filtering, and writes a CSV with columns: doc_id, sentence_blurb.

Usage:
    python docx_to_csv.py --input /path/to/docx/folder --output output.csv
    python docx_to_csv.py --input /path/to/docx/folder --output output.csv --blurb-size 2
    python docx_to_csv.py --input /path/to/docx/folder --output output.csv --no-filter

Requirements:
    pip install python-docx
    (pandoc optional but improves extraction quality — install from https://pandoc.org)
"""

import os
import re
import csv
import argparse
import subprocess
import tempfile
import logging
from pathlib import Path

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text_pandoc(docx_path: str) -> str:
    """Extract plain text via pandoc (preferred — handles headers, tables, etc.)"""
    try:
        result = subprocess.run(
            ["pandoc", "--track-changes=all", "-t", "plain", docx_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def extract_text_python_docx(docx_path: str) -> str:
    """Fallback: extract text using python-docx."""
    try:
        from docx import Document
        doc = Document(docx_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text.strip())
        return "\n".join(p for p in paragraphs if p)
    except Exception as e:
        log.warning(f"python-docx failed on {docx_path}: {e}")
        return ""


def extract_text(docx_path: str) -> str:
    text = extract_text_pandoc(docx_path)
    if text is None:
        log.debug(f"pandoc unavailable, falling back to python-docx for {docx_path}")
        text = extract_text_python_docx(docx_path)
    return text or ""


# ---------------------------------------------------------------------------
# Sentence splitting
# ---------------------------------------------------------------------------

def split_sentences(text: str) -> list[str]:
    """
    Split text into sentences using a simple but robust regex approach.
    Handles common abbreviations, decimals, and multi-sentence patterns.
    """
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text).strip()

    # Split on sentence-ending punctuation followed by whitespace + capital letter
    # This pattern handles: ". " "! " "? " but avoids splitting on abbreviations like U.S.
    sentence_endings = re.compile(
        r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\!|\?)\s+(?=[A-Z"])'
    )
    raw = sentence_endings.split(text)

    # Clean up each sentence
    sentences = []
    for s in raw:
        s = s.strip()
        if s:
            sentences.append(s)

    return sentences


# ---------------------------------------------------------------------------
# Filtering
# ---------------------------------------------------------------------------

# Patterns that strongly indicate boilerplate / irrelevant content
BOILERPLATE_PATTERNS = [
    re.compile(r'^\s*page\s+\d+\s*(of\s+\d+)?\s*$', re.IGNORECASE),          # "Page 3 of 12"
    re.compile(r'^\s*\d+\s*$'),                                                 # lone numbers
    re.compile(r'^\s*(section|chapter|part|exhibit|appendix)\s+[IVXLC\d]+\s*[:\-]?\s*$', re.IGNORECASE),  # "Section IV:"
    re.compile(r'^\s*table\s+of\s+contents?\s*$', re.IGNORECASE),
    re.compile(r'^\s*confidential\b', re.IGNORECASE),
    re.compile(r'^\s*this\s+document\s+is\s+(confidential|proprietary)', re.IGNORECASE),
    re.compile(r'^\s*all\s+rights?\s+reserved', re.IGNORECASE),
    re.compile(r'^\s*(signed|dated|acknowledged)\s+(this|on)\b', re.IGNORECASE),
    re.compile(r'^\s*\[.*?\]\s*$'),                                             # [REDACTED] etc.
    re.compile(r'^\s*_{3,}\s*$'),                                               # signature lines ___
    re.compile(r'^\s*[-=*]{3,}\s*$'),                                           # dividers --- ===
    re.compile(r'^\s*\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}\s*$'),                 # bare dates
    re.compile(r'^\s*(exhibit|attachment|enclosure)\s+[A-Z\d]+\s*$', re.IGNORECASE),
    # LexisNexis / Westlaw headnote boilerplate
    re.compile(r'HN\d+\['),                                                     # headnote refs HN1[], HN2[]
    re.compile(r'LexisNexis'),                                                  # LexisNexis® branding
    re.compile(r'^\s*(Civil Procedure|Constitutional Law|Criminal Law|Governments|Evidence|Immigration Law|Torts|Contracts|Labor|Healthcare)\s*[>|]', re.IGNORECASE),  # taxonomy paths
    re.compile(r'\bCore Terms\b', re.IGNORECASE),                               # LexisNexis Core Terms header
    re.compile(r'\bLexisNexis®?\s+Headnotes\b', re.IGNORECASE),
    re.compile(r'Reporter\s+\d+\s+[A-Z]\.?\d*', re.IGNORECASE),               # Reporter citation blocks
    re.compile(r'^\s*\d+\s*\*+\s*$'),                                           # page break markers like *264*
    re.compile(r'U\.S\.\s+LEXIS\s+\d+', re.IGNORECASE),                        # bare LEXIS citation lines
    re.compile(r'^\s*\[\*+\d*\]\s*$'),                                          # [**1] paragraph markers
]

MIN_WORD_COUNT = 6          # drop blurbs with fewer words
MIN_ALPHA_RATIO = 0.4       # at least 40% alphabetic characters
MAX_CAPS_RATIO = 0.7        # drop if more than 70% of letters are uppercase (likely a header)


def is_relevant(blurb: str) -> bool:
    """Return True if the blurb appears substantive and not boilerplate."""
    stripped = blurb.strip()

    # Too short
    words = stripped.split()
    if len(words) < MIN_WORD_COUNT:
        return False

    # Boilerplate regex matches
    for pattern in BOILERPLATE_PATTERNS:
        if pattern.search(stripped):
            return False

    # Not enough alphabetic content (numbers/symbols heavy)
    alpha_chars = sum(c.isalpha() for c in stripped)
    total_chars = len(stripped)
    if total_chars > 0 and (alpha_chars / total_chars) < MIN_ALPHA_RATIO:
        return False

    # All-caps or nearly all-caps = likely a section header
    if alpha_chars > 0 and (sum(c.isupper() for c in stripped) / alpha_chars) > MAX_CAPS_RATIO:
        return False

    return True


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def chunk_sentences(sentences: list[str], blurb_size: int = 2) -> list[str]:
    """
    Group sentences into overlapping blurbs of `blurb_size` sentences.
    Uses a sliding window with step=1 for maximum context coverage.
    """
    if not sentences:
        return []
    blurbs = []
    for i in range(len(sentences)):
        chunk = sentences[i: i + blurb_size]
        blurb = " ".join(chunk).strip()
        if blurb:
            blurbs.append(blurb)
    return blurbs


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Human-smuggling folder defaults
# ---------------------------------------------------------------------------

HUMAN_SMUGGLING_FOLDER = (
    "/Users/eli/projects/humansmug/humansmug/Additional Cases/(human smuggling)"
)
DEFAULT_OUTPUT_CSV = (
    "/Users/eli/projects/humansmug/humansmug/Additional Cases/(human smuggling)/Results/output.csv"
)

# Subfolders to skip entirely (contain outputs, not source docs)
SKIP_SUBFOLDERS = {"Results"}

# Filename patterns to skip (index / TOC files, not actual cases)
SKIP_FILENAME_PATTERNS = [
    re.compile(r'doclist', re.IGNORECASE),
]


def _is_skipped(path: Path, root: Path) -> bool:
    """Return True if this path lives inside a skipped subfolder or matches a skip pattern."""
    relative_parts = path.relative_to(root).parts
    if any(part in SKIP_SUBFOLDERS for part in relative_parts):
        return True
    for pattern in SKIP_FILENAME_PATTERNS:
        if pattern.search(path.name):
            return True
    return False


def process_folder(
    input_folder: str,
    output_csv: str,
    blurb_size: int = 2,
    apply_filter: bool = True,
) -> None:
    input_path = Path(input_folder)

    # Collect .docx files case-insensitively (handles both .docx and .DOCX)
    all_files = sorted(
        f for f in input_path.rglob("*")
        if f.is_file() and f.suffix.lower() == ".docx"
    )
    docx_files = [f for f in all_files if not _is_skipped(f, input_path)]

    if not docx_files:
        log.error(f"No .docx files found in {input_folder}")
        return

    log.info(f"Found {len(docx_files)} .docx files")

    total_rows = 0
    filtered_rows = 0

    with open(output_csv, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile, quoting=csv.QUOTE_ALL)
        writer.writerow(["doc_id", "sentence_blurb"])

        for i, docx_path in enumerate(docx_files, 1):
            # Include the immediate batch subfolder (e.g. "1-100") for provenance
            relative = docx_path.relative_to(input_path)
            batch = relative.parts[0] if len(relative.parts) > 1 else ""
            doc_id = f"{batch}/{docx_path.stem}" if batch else docx_path.stem
            log.info(f"[{i}/{len(docx_files)}] Processing: {doc_id}")

            try:
                text = extract_text(str(docx_path))
                if not text.strip():
                    log.warning(f"  No text extracted from {doc_id}")
                    continue

                sentences = split_sentences(text)
                blurbs = chunk_sentences(sentences, blurb_size=blurb_size)

                for blurb in blurbs:
                    total_rows += 1
                    if apply_filter and not is_relevant(blurb):
                        filtered_rows += 1
                        continue
                    writer.writerow([doc_id, blurb])

            except Exception as e:
                log.error(f"  Failed on {doc_id}: {e}")

    kept = total_rows - filtered_rows
    log.info(f"\nDone. {kept} blurbs written to {output_csv}")
    if apply_filter:
        log.info(f"Filtered out {filtered_rows} / {total_rows} blurbs ({100*filtered_rows//max(total_rows,1)}%)")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract docx text into a sentence-blurb CSV.")
    parser.add_argument("--input",       default=HUMAN_SMUGGLING_FOLDER,
                        help=f"Folder containing .docx files (default: human-smuggling folder)")
    parser.add_argument("--output",      default=DEFAULT_OUTPUT_CSV,
                        help=f"Output CSV file path (default: Results/output.csv inside input folder)")
    parser.add_argument("--blurb-size",  type=int, default=2, choices=[1, 2, 3],
                        help="Number of sentences per blurb (default: 2)")
    parser.add_argument("--no-filter",   action="store_true",
                        help="Disable rule-based filtering (keep all blurbs)")
    args = parser.parse_args()

    # Ensure the output directory exists
    Path(args.output).parent.mkdir(parents=True, exist_ok=True)

    process_folder(
        input_folder=args.input,
        output_csv=args.output,
        blurb_size=args.blurb_size,
        apply_filter=not args.no_filter,
    )